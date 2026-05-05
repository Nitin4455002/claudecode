#!/usr/bin/env python3
"""Build a polished DOCX from email-sequence-rankframe.md."""
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


SRC = "/home/user/claudecode/email-sequence-rankframe.md"
OUT = "/home/user/claudecode/RankFrame-Email-Sequence.docx"

PURPLE = RGBColor(0x6B, 0x46, 0xC1)
DARK = RGBColor(0x1F, 0x1F, 0x1F)
GRAY = RGBColor(0x6B, 0x72, 0x80)
LIGHT_GRAY = RGBColor(0xE5, 0xE7, 0xEB)


def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def add_hr(doc):
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def style_doc(doc):
    styles = doc.styles
    normal = styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)
    normal.font.color.rgb = DARK

    for level, size, color in [
        ('Heading 1', 24, PURPLE),
        ('Heading 2', 18, PURPLE),
        ('Heading 3', 14, DARK),
        ('Heading 4', 12, DARK),
    ]:
        s = styles[level]
        s.font.name = 'Calibri'
        s.font.size = Pt(size)
        s.font.color.rgb = color
        s.font.bold = True


def add_inline(paragraph, text):
    """Render bold (**...**) and inline code (`...`) in a paragraph run."""
    parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            r = paragraph.add_run(part[2:-2])
            r.bold = True
        elif part.startswith('`') and part.endswith('`'):
            r = paragraph.add_run(part[1:-1])
            r.font.name = 'Consolas'
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0xB9, 0x18, 0x4C)
        else:
            paragraph.add_run(part)


def parse_table(lines, i):
    """Parse a markdown table starting at index i. Returns (rows, new_i)."""
    rows = []
    while i < len(lines) and '|' in lines[i]:
        line = lines[i].strip()
        if re.match(r'^\|?\s*[-:|\s]+\|?\s*$', line):
            i += 1
            continue
        cells = [c.strip() for c in line.strip('|').split('|')]
        rows.append(cells)
        i += 1
    return rows, i


def add_table(doc, rows):
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Light Grid Accent 1'
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            if c_idx >= len(table.rows[r_idx].cells):
                continue
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            add_inline(p, cell_text)
            if r_idx == 0:
                set_cell_bg(cell, '6B46C1')
                for run in p.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.size = Pt(10)
            else:
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()


def add_code_block(doc, code_lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    # background shading on the paragraph
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F3F4F6')
    p_pr.append(shd)
    r = p.add_run('\n'.join(code_lines))
    r.font.name = 'Consolas'
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)


def build():
    doc = Document()
    style_doc(doc)

    # Page margins
    for section in doc.sections:
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)

    # Cover block
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = title.add_run('RankFrame Email Sequence')
    r.font.size = Pt(28)
    r.font.bold = True
    r.font.color.rgb = PURPLE

    subtitle = doc.add_paragraph()
    r = subtitle.add_run('Trial to Paid · Zapier + MailerLite Integration Spec')
    r.font.size = Pt(14)
    r.font.color.rgb = GRAY

    meta = doc.add_paragraph()
    r = meta.add_run('The first and only native SEO plugin for Framer.')
    r.font.size = Pt(11)
    r.italic = True
    r.font.color.rgb = GRAY

    add_hr(doc)

    # Read source
    with open(SRC, 'r') as f:
        lines = f.read().split('\n')

    # Skip the original H1 + first separator (we already wrote a cover)
    # Find the first '---' after the title block and start from there.
    skip_until = 0
    seen_h1 = False
    for idx, line in enumerate(lines):
        if line.startswith('# ') and not seen_h1:
            seen_h1 = True
            continue
        if seen_h1 and line.strip() == '---':
            skip_until = idx + 1
            break
    lines = lines[skip_until:]

    i = 0
    in_code = False
    code_buf = []

    while i < len(lines):
        line = lines[i]

        # Code fence
        if line.strip().startswith('```'):
            if in_code:
                add_code_block(doc, code_buf)
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        stripped = line.strip()

        # HR
        if stripped == '---':
            add_hr(doc)
            i += 1
            continue

        # Headings
        if stripped.startswith('#### '):
            doc.add_heading(stripped[5:], level=4)
            i += 1
            continue
        if stripped.startswith('### '):
            doc.add_heading(stripped[4:], level=3)
            i += 1
            continue
        if stripped.startswith('## '):
            doc.add_heading(stripped[3:], level=2)
            i += 1
            continue
        if stripped.startswith('# '):
            doc.add_heading(stripped[2:], level=1)
            i += 1
            continue

        # Tables
        if '|' in line and i + 1 < len(lines) and re.match(r'^\|?\s*[-:|\s]+\|?\s*$', lines[i+1].strip()):
            rows, i = parse_table(lines, i)
            add_table(doc, rows)
            continue

        # Blockquote
        if stripped.startswith('> '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.6)
            p_pr = p._p.get_or_add_pPr()
            p_bdr = OxmlElement('w:pBdr')
            left = OxmlElement('w:left')
            left.set(qn('w:val'), 'single')
            left.set(qn('w:sz'), '18')
            left.set(qn('w:space'), '6')
            left.set(qn('w:color'), '6B46C1')
            p_bdr.append(left)
            p_pr.append(p_bdr)
            add_inline(p, stripped[2:])
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = GRAY
            i += 1
            continue

        # Bullet list
        m = re.match(r'^(\s*)-\s+(.*)$', line)
        if m:
            indent_spaces = len(m.group(1))
            level = indent_spaces // 2
            text = m.group(2)
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Cm(0.6 + level * 0.6)
            # Checkbox lists "[ ]"
            if text.startswith('[ ] '):
                text = '☐ ' + text[4:]
            elif text.startswith('[x] '):
                text = '☑ ' + text[4:]
            add_inline(p, text)
            i += 1
            continue

        # Numbered list
        m = re.match(r'^(\s*)(\d+)\.\s+(.*)$', line)
        if m:
            text = m.group(3)
            p = doc.add_paragraph(style='List Number')
            add_inline(p, text)
            i += 1
            continue

        # Italic-only line wrapped in *...*
        if stripped.startswith('*') and stripped.endswith('*') and not stripped.startswith('**'):
            p = doc.add_paragraph()
            r = p.add_run(stripped.strip('*'))
            r.italic = True
            r.font.color.rgb = GRAY
            r.font.size = Pt(10)
            i += 1
            continue

        # Empty line
        if stripped == '':
            i += 1
            continue

        # Default paragraph
        p = doc.add_paragraph()
        add_inline(p, line)
        i += 1

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == '__main__':
    build()
